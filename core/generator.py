# -*- coding: utf-8 -*-
"""Geração de documentos Word — compartilhado por Notas e Reajuste."""
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.config import (
    ALIQUOTAS, CIDADE_DISPLAY, SERVICE_LABELS,
    CONTRATO_HEADER, REAJ_HEADER,
    INSS_ISENTO, BASE_CALC, ADM_LABEL,
    MED_PARCIAL, MED_REAJ, TOTAL_MED,
)
from core.extractor import fmt


def gerar_texto_medicao(info: dict, nome_upper: str, vals: dict) -> str:
    cidade = CIDADE_DISPLAY.get(nome_upper.rstrip(), nome_upper)
    linhas = [CONTRATO_HEADER.format(cidade=cidade), "", BASE_CALC]
    base = 0.0
    for key, aliq in [("pav", ALIQUOTAS["pav"]), ("canteiro", ALIQUOTAS["canteiro"]),
                      ("rocada", ALIQUOTAS["rocada"]), ("terra", ALIQUOTAS["terra"])]:
        if vals[key] > 0:
            b = vals[key] * aliq
            base += b
            linhas.append(f"{SERVICE_LABELS[key]}: {fmt(vals[key])} x {int(aliq*100)}% = R$ {fmt(b)}")
    if base > 0:
        linhas.append(f"Valor do INSS: R$ {fmt(base)} x 11% R$ {fmt(base * ALIQUOTAS['inss'])}")
    if vals["adm"] > 0:
        linhas.append(f"{ADM_LABEL}: {fmt(vals['adm'])} {INSS_ISENTO}")
    linhas += [
        "",
        MED_PARCIAL.format(num=info["num_medicao"], periodo=info["periodo"], total=fmt(vals["total"])),
        TOTAL_MED.format(total=fmt(vals["total"])),
    ]
    return "\n".join(linhas)


def gerar_texto_reajuste(info: dict, nome: str, reaj: dict) -> str:
    cidade = CIDADE_DISPLAY.get(nome.strip(), nome)
    linhas = [REAJ_HEADER.format(cidade=cidade), "", BASE_CALC]
    base = 0.0
    for key, aliq in [("pav", ALIQUOTAS["pav"]), ("canteiro", ALIQUOTAS["canteiro"]),
                      ("rocada", ALIQUOTAS["rocada"]), ("terra", ALIQUOTAS["terra"])]:
        if reaj[key] > 0:
            b = reaj[key] * aliq
            base += b
            linhas.append(f"{SERVICE_LABELS[key]}: {fmt(reaj[key])} x {int(aliq*100)}% = R$ {fmt(b)}")
    if base > 0:
        linhas.append(f"Valor do INSS: R$ {fmt(base)} x 11% R$ {fmt(base * ALIQUOTAS['inss'])}")
    if reaj["adm"] > 0:
        linhas.append(f"{ADM_LABEL}: {fmt(reaj['adm'])} {INSS_ISENTO}")
    linhas += [
        "",
        MED_REAJ.format(num=info["num_medicao"], periodo=info["periodo"], total=fmt(reaj["total"])),
        TOTAL_MED.format(total=fmt(reaj["total"])),
    ]
    return "\n".join(linhas)


def gerar_word_medicao(info: dict, cidades: dict) -> bytes:
    """Retorna bytes do .docx de medição."""
    doc = Document()
    primeiro = True
    for nome_upper, vals in cidades.items():
        if not primeiro:
            doc.add_paragraph()
        cidade = CIDADE_DISPLAY.get(nome_upper.rstrip(), nome_upper)
        p = doc.add_paragraph(CONTRATO_HEADER.format(cidade=cidade))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
        doc.add_paragraph(BASE_CALC)
        base = 0.0
        for key, aliq in [("pav", ALIQUOTAS["pav"]), ("canteiro", ALIQUOTAS["canteiro"]),
                           ("rocada", ALIQUOTAS["rocada"]), ("terra", ALIQUOTAS["terra"])]:
            if vals[key] > 0:
                b = vals[key] * aliq
                base += b
                doc.add_paragraph(f"{SERVICE_LABELS[key]}: {fmt(vals[key])} x {int(aliq*100)}% = R$ {fmt(b)}")
        if base > 0:
            doc.add_paragraph(f"Valor do INSS: R$ {fmt(base)} x 11% R$ {fmt(base * ALIQUOTAS['inss'])}")
        if vals["adm"] > 0:
            doc.add_paragraph(f"{ADM_LABEL}: {fmt(vals['adm'])} {INSS_ISENTO}")
        doc.add_paragraph()
        doc.add_paragraph(MED_PARCIAL.format(
            num=info["num_medicao"], periodo=info["periodo"], total=fmt(vals["total"])))
        doc.add_paragraph(TOTAL_MED.format(total=fmt(vals["total"])))
        primeiro = False
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def gerar_word_reajuste(info: dict, reajuste: dict) -> bytes:
    """Retorna bytes do .docx de reajustamento."""
    doc = Document()
    primeiro = True
    for nome, reaj in reajuste.items():
        if not primeiro:
            doc.add_paragraph()
        cidade = CIDADE_DISPLAY.get(nome.strip(), nome)
        p = doc.add_paragraph(REAJ_HEADER.format(cidade=cidade))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
        doc.add_paragraph(BASE_CALC)
        base = 0.0
        for key, aliq in [("pav", ALIQUOTAS["pav"]), ("canteiro", ALIQUOTAS["canteiro"]),
                           ("rocada", ALIQUOTAS["rocada"]), ("terra", ALIQUOTAS["terra"])]:
            if reaj[key] > 0:
                b = reaj[key] * aliq
                base += b
                doc.add_paragraph(f"{SERVICE_LABELS[key]}: {fmt(reaj[key])} x {int(aliq*100)}% = R$ {fmt(b)}")
        if base > 0:
            doc.add_paragraph(f"Valor do INSS: R$ {fmt(base)} x 11% R$ {fmt(base * ALIQUOTAS['inss'])}")
        if reaj["adm"] > 0:
            doc.add_paragraph(f"{ADM_LABEL}: {fmt(reaj['adm'])} {INSS_ISENTO}")
        doc.add_paragraph()
        doc.add_paragraph(MED_REAJ.format(
            num=info["num_medicao"], periodo=info["periodo"], total=fmt(reaj["total"])))
        doc.add_paragraph(TOTAL_MED.format(total=fmt(reaj["total"])))
        primeiro = False
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
